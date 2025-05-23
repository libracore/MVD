# -*- coding: utf-8 -*-
# Copyright (c) 2023, libracore and contributors
# For license information, please see license.txt

from __future__ import unicode_literals
import frappe
from frappe.model.document import Document
from datetime import date, timedelta
from frappe.utils.data import getdate
from frappe import _

class ArbeitsplanBeratung(Document):
    def before_save(self):
        if not self.titel or self.titel == '':
            self.titel = "{0} - {1} ({2})".format(frappe.utils.get_datetime(self.from_date).strftime('%d.%m.%Y'), \
            frappe.utils.get_datetime(self.to_date).strftime('%d.%m.%Y'), self.sektion_id)
    
    def validate(self):
        self.validate_date()
        self.validate_overlapping()
    
    def validate_date(self):
        if self.to_date < self.from_date:
            frappe.throw("Das Bis-Datum muss nach dem Von-Datum sein.")
    
    def validate_overlapping(self):
        overlapping = frappe.db.sql("""SELECT
                                            `name`
                                        FROM `tabArbeitsplan Beratung`
                                        WHERE `sektion_id` = '{sektion}'
                                        AND `name` != '{himself}'
                                        AND (
                                            `from_date` BETWEEN '{from_date}' AND '{to_date}'
                                            OR
                                            `to_date` BETWEEN '{from_date}' AND '{to_date}'
                                        )
                                        AND `docstatus` != 2
                                        """.format(sektion=self.sektion_id, \
                                                    from_date=self.from_date, \
                                                    to_date=self.to_date, \
                                                    himself=self.name), as_list=True)
        if len(overlapping) > 0:
            overlaps = ''
            for overlap in overlapping:
                overlaps += overlap[0] + "<br>"
            frappe.throw("Der Datumsbereich dieses Arbeitsplan überschneited sich mit:<br>{0}".format(overlaps))

    def get_personen(self, einzel, person):
        def get_beratungspersonen(weekday, sektion):
            person_query = ''
            if einzel:
                bereits_erfasst = False
                for beratungsperson in self.einteilung:
                    if beratungsperson.beratungsperson == person:
                        frappe.throw("""Die Einteilungen für {0} wurden dem Arbeitsplan bereits hinzugefügt.""".format(person))
                
                person_query = """AND `as`.`parent` = '{0}'""".format(person)
                if not person:
                    frappe.throw("Bitte wählen Sie ein*e Berater*in aus")
            
            weekday_query = [
                "Montag",
                "Dienstag",
                "Mittwoch",
                "Donnerstag",
                "Freitag",
                "Samstag",
                "Sonntag"
            ]
            beratungspersonen = frappe.db.sql("""SELECT
                                                    `as`.`parent` AS `beratungsperson`,
                                                    `as`.`from` AS `from_time`,
                                                    `as`.`to` AS `to_time`,
                                                    `as`.`ort` AS `art_ort`,
                                                    `tk`.`dispositions_hinweis`
                                                from `tabArbeitsplan Standardzeit` AS `as`
                                                LEFT JOIN `tabTermin Kontaktperson` AS `tk` ON `as`.`parent` = `tk`.`name`
                                                WHERE `tk`.`sektion_id` = '{sektion}'
                                                AND `tk`.`disabled` != 1
                                                AND `as`.`day` = '{weekday}'
                                                {person_query}""".format(weekday=weekday_query[weekday], \
                                                                    sektion=sektion, person_query=person_query), as_dict=True)
            if len(beratungspersonen) > 0:
                return beratungspersonen
            else:
                return False
        
        start_date = getdate(self.from_date)
        end_date = getdate(self.to_date)
        delta = timedelta(days=1)
        einteilung_list = []
        dispositions_hinweise = {}
        dispositions_hinweise_txt = ""

        while start_date <= end_date:
            beratungspersonen = get_beratungspersonen(start_date.weekday(), self.sektion_id)
            if beratungspersonen:
                for beratungsperson in beratungspersonen:
                    x = {}
                    x['date'] = start_date.strftime("%Y-%m-%d")
                    x['from_time'] = beratungsperson.from_time
                    x['to_time'] = beratungsperson.to_time
                    x['art_ort'] = beratungsperson.art_ort
                    x['beratungsperson'] = beratungsperson.beratungsperson
                    einteilung_list.append(x)
                    if beratungsperson.beratungsperson not in dispositions_hinweise:
                        if beratungsperson.dispositions_hinweis:
                            dispositions_hinweise[beratungsperson.beratungsperson] = beratungsperson.dispositions_hinweis
            start_date += delta
        
        sorted_einteilung_list = sorted(einteilung_list, key = lambda x: (x['art_ort'], x['date'], x['from_time']))

        for eintrag in sorted_einteilung_list:
            row = self.append('einteilung', {})
            row.date = eintrag['date']
            row.from_time = eintrag['from_time']
            row.to_time = eintrag['to_time']
            row.art_ort = eintrag['art_ort']
            row.beratungsperson = eintrag['beratungsperson']
        
        for key in dispositions_hinweise:
            dispositions_hinweise_txt += "{0}: {1}\n".format(key, dispositions_hinweise[key])
        self.dispositions_hinweis = dispositions_hinweise_txt
        
        self.save()
    
    def order_einteilung(self, column, order):
        einteilung_list = self.einteilung
        sorted_einteilung_list = sorted(einteilung_list, key = lambda x: x.get(column), reverse=False if order == 'down' else True)

        loop = 1
        for entry in sorted_einteilung_list:
            entry.idx = loop
            loop += 1
        
        self.einteilung = sorted_einteilung_list
        self.save()
        return

@frappe.whitelist()
def zeige_verfuegbarkeiten(sektion, datum, beraterin=None, ort=None, marked=None, short_results=1, art=None):
    von_datum = getdate(datum)
    delta = timedelta(days=1)
    if int(short_results) == 1:
        bis_datum = von_datum + timedelta(days=14)
    else:
        bis_datum = von_datum + timedelta(days=365)
    if marked:
        marked = marked.split("-")
    beraterin_filter = ''
    ort_filter = ''
    art_filter = ''
    verfuegbarkeiten_html = ""
    if beraterin and beraterin != '':
        beraterin_filter = '''AND `beratungsperson` = '{0}' '''.format(beraterin)
    if ort and ort != '' and ort != ' ':
        ort_filter = '''AND `art_ort` = '{0}' '''.format(ort)
    if art and art != 'telefonisch':
        art_filter = """AND `art_ort` NOT LIKE 'Telefon%'"""

    
    while von_datum <= bis_datum:
        if von_datum.strftime('%A') not in ['Saturday', 'Sunday']:
            zugeteilte_beratungspersonen = frappe.db.sql("""
                                                            SELECT *
                                                            FROM `tabAPB Zuweisung`
                                                            WHERE `date` = '{von_datum}'
                                                            AND `name` NOT IN (
                                                                SELECT `abp_referenz`
                                                                FROM `tabBeratung Termin`
                                                                WHERE `abp_referenz` IS NOT NULL
                                                            )
                                                            AND `parent` IN (
                                                                SELECT `name`
                                                                FROM `tabArbeitsplan Beratung`
                                                                WHERE `sektion_id` = '{sektion}'
                                                            )
                                                            {beraterin_filter}
                                                            {ort_filter}
                                                            {art_filter}
                                                        """.format(von_datum=von_datum.strftime("%Y-%m-%d"), \
                                                                   beraterin_filter=beraterin_filter, ort_filter=ort_filter, \
                                                                    art_filter=art_filter, sektion=sektion), as_dict=True)
            if int(short_results) == 1:
                verfuegbarkeiten_html += """
                    <p style="margin-bottom: 0px !important;"><b>{wochentag}, {datum}</b></p>
                """.format(wochentag=_(von_datum.strftime('%A')), datum=von_datum.strftime('%d.%m.%y'))
            else:
                if len(zugeteilte_beratungspersonen) > 0:
                    verfuegbarkeiten_html += """
                    <p style="margin-bottom: 0px !important;"><b>{wochentag}, {datum}</b></p>
                """.format(wochentag=_(von_datum.strftime('%A')), datum=von_datum.strftime('%d.%m.%y'))
            
            if len(zugeteilte_beratungspersonen) > 0:
                zugeteilte_beratungspersonen_liste = []
                for zugeteilte_beratungsperson in zugeteilte_beratungspersonen:
                    x = {
                        'ort': zugeteilte_beratungsperson.art_ort.replace("({sektion})".format(sektion=sektion), ""),
                        'ort_mit_sektion': zugeteilte_beratungsperson.art_ort,
                        'from_time': zugeteilte_beratungsperson.from_time,
                        'to_time': zugeteilte_beratungsperson.to_time,
                        'beratungsperson': zugeteilte_beratungsperson.beratungsperson.replace("({sektion})".format(sektion=sektion), ""),
                        'beratungsperson_mit_sektion': zugeteilte_beratungsperson.beratungsperson,
                        'name': zugeteilte_beratungsperson.name
                    }
                    zugeteilte_beratungspersonen_liste.append(x)
                sorted_list = sorted(zugeteilte_beratungspersonen_liste, key = lambda x: (x['ort'], x['from_time'], x['beratungsperson']))
                for entry in sorted_list:
                    checked = ''
                    if entry['name'] in marked:
                        checked = 'checked'
                    verfuegbarkeiten_html += """
                        <div class="form-group frappe-control input-max-width" style="margin-bottom: 0px !important;" data-fieldtype="Check">
                            <div class="checkbox" style="margin-bottom: 0px !important; margin-top: 0px !important;">
                                <label>
                                    <span class="input-area"><input type="checkbox" {checked} autocomplete="off" class="input-with-feedback" data-fieldtype="Check" data-abpzuweisung="{abpzuweisung}" data-ort="{ort_mit_sektion}" data-beratungsperson="{beratungsperson_mit_sektion}" onclick="cur_dialog.checkbox_clicked(this);"></span>
                                    <span class="label-area small">{from_time} - {to_time} / {ort} / {beratungsperson}</span>
                                </label>
                            </div>
                        </div>
                    """.format(from_time=':'.join(str(entry['from_time']).split(':')[:2]), \
                            to_time=':'.join(str(entry['to_time']).split(':')[:2]), \
                                ort=entry['ort'], beratungsperson=entry['beratungsperson'], \
                                abpzuweisung=entry['name'], ort_mit_sektion=entry['ort_mit_sektion'], \
                                checked=checked, beratungsperson_mit_sektion=entry['beratungsperson_mit_sektion'])
            else:
                if int(short_results) == 1:
                    verfuegbarkeiten_html += """<pstyle="margin-top: 0px !important;">Kein(e) Berater*in verfügbar</p>"""
        von_datum += delta
    
    return verfuegbarkeiten_html

@frappe.whitelist()
def verwendete_einteilungen(arbeitsplan_beratung):
    verwendet = frappe.db.sql("""
        SELECT *
        FROM `tabBeratung Termin`
        WHERE `abp_referenz` IN (
            SELECT `name` FROM `tabAPB Zuweisung`
            WHERE `parent` = '{arbeitsplan_beratung}'
        )
    """.format(arbeitsplan_beratung=arbeitsplan_beratung), as_dict=True)

    verwendet_as_sting = ''
    reset_values = []

    for v in verwendet:
        verwendet_as_sting += "-{0}".format(v.abp_referenz)

        values_for_reset = frappe.db.sql("""
            SELECT *
            FROM `tabAPB Zuweisung`
            WHERE `name`  = '{zuweisung}'
        """.format(zuweisung=v.abp_referenz), as_dict=True)

        reset_values.append({
            'referenz': v.abp_referenz,
            'reset_data': [
                values_for_reset[0].art_ort,
                values_for_reset[0].date,
                values_for_reset[0].from_time,
                values_for_reset[0].to_time,
                values_for_reset[0].beratungsperson
            ]
        })
    
    return {
        'einteilung_verwendet': verwendet_as_sting,
        'reset_values': reset_values
    }

@frappe.whitelist()
def get_termin_uebersicht(berater_in, von=None, bis=None):
    von_filter = ''
    bis_filter = ''
    if von:
        von_filter = """AND `bt`.`von` >= '{0} 00:00:00'""".format(von)
    if bis:
        bis_filter = """AND `bt`.`bis` <= '{0} 23:59:59'""".format(bis)

    berater_in = get_berater_in_from_hash(berater_in)

    termine = frappe.db.sql("""
        SELECT
            `bt`.`von`,
            `bt`.`bis`,
            `bt`.`art`,
            `bt`.`ort`,
            `bt`.`telefonnummer`,
            REPLACE(`mitgl`.`adressblock`, '\n', '<br>') AS `adressblock`,
            `ber`.`sektion_id`,
            `mitgl`.`mitglied_nr`,
            `mitgl`.`eintrittsdatum`,
            `mitgl`.`bezahltes_mitgliedschaftsjahr`,
            `bt`.`creation`,
            `bt`.`owner`,
            `bt`.`notiz`,
            '' AS `von_zeit`,
            `ber`.`name` AS `beratung_name`,
            `ber`.`beratungskategorie` AS `beratungskategorie`,
            `ber`.`beratungskategorie_2` AS `beratungskategorie_2`,
            `ber`.`beratungskategorie_3` AS `beratungskategorie_3`,
            `bt`.`abp_referenz` AS `abp_referenz`,
            NULL AS `von_bis_str`,
            NULL AS `hat_attachement`,
            NULL AS `created_by`
        FROM `tabBeratung Termin` AS `bt`
        LEFT JOIN `tabBeratung` AS `ber` ON `bt`.`parent` = `ber`.`name`
        LEFT JOIN `tabMitgliedschaft` AS `mitgl` ON `ber`.`mv_mitgliedschaft` = `mitgl`.`name`
        WHERE `berater_in` = '{berater_in}'
        {von_filter}
        {bis_filter}
        ORDER BY `von` ASC
    """.format(berater_in=berater_in, von_filter=von_filter, bis_filter=bis_filter), as_dict=True)

    vergebene_termin_liste = []
    for termin in termine:
        if termin.abp_referenz:
            termin.hat_attachement = 1 if frappe.db.sql("""SELECT COUNT(`name`) AS `qty` FROM `tabBeratungsdateien` WHERE `parent` = '{termin}'""".format(termin=termin.beratung_name), as_dict=True)[0].qty > 0 else 0
            termin.created_by = get_created_by(termin.owner)
            vergebene_termin_liste.append(termin.abp_referenz)
    
    freie_termine = get_freie_termine(vergebene_termin_liste, von, bis, berater_in)
    for freier_termin in freie_termine:
        termine.append(freier_termin)
    
    sortierte_termine = sorted(termine, key=lambda d: d['von'])
    return sortierte_termine

def get_freie_termine(vergebene_termin_liste, von, bis, berater_in):
    von_filter = ''
    bis_filter = ''
    if von:
        von_filter = """AND `date` >= '{0}'""".format(von)
    if bis:
        bis_filter = """AND `date` <= '{0}'""".format(bis)
    
    freie_termine = frappe.db.sql("""
                                  SELECT
                                    CONCAT(`date`, ' ', `from_time`) AS `von`,
                                    CONCAT(`date`, ' ', `to_time`) AS `bis`,
                                    `art_ort` AS `ort`,
                                    `beratungsperson` AS `berater_in`,
                                    NULL AS `von_bis_str`
                                  FROM `tabAPB Zuweisung`
                                  WHERE `name` NOT IN ('{0}')
                                  AND `beratungsperson` = '{1}'
                                  {2}
                                  {3}
                                  """.format("', '".join(vergebene_termin_liste), berater_in, von_filter, bis_filter), as_dict=True)
    for freier_termin in freie_termine:
        freier_termin.von = frappe.utils.get_datetime(freier_termin.von)
        freier_termin.bis = frappe.utils.get_datetime(freier_termin.bis)
    return freie_termine

@frappe.whitelist()
def get_arbeitsplan_pdf(berater_in, von=None, bis=None):
    termine = get_termin_uebersicht(berater_in, von, bis)
    for termin in termine:
        new_von_zeit = "{0}:{1}".format("{0}".format(termin.von).split(" ")[1].split(":")[0], "{0}".format(termin.von).split(" ")[1].split(":")[1])
        new_von = "{0}.{1}.{2}".format("{0}".format(termin.von).split(" ")[0].split("-")[2], "{0}".format(termin.von).split(" ")[0].split("-")[1], "{0}".format(termin.von).split(" ")[0].split("-")[0])
        new_bis = "{0}:{1}".format("{0}".format(termin.bis).split(" ")[1].split(":")[0], "{0}".format(termin.bis).split(" ")[1].split(":")[1])
        termin.von_zeit = new_von_zeit
        termin.von = new_von
        termin.bis = new_bis
        if termin.abp_referenz:
            termin.eintrittsdatum = frappe.utils.getdate(termin.eintrittsdatum).strftime("%d.%m.%Y")
            termin.creation = frappe.utils.getdate(termin.creation).strftime("%d.%m.%Y")
            termin.hat_attachement = 1 if frappe.db.sql("""SELECT COUNT(`name`) AS `qty` FROM `tabBeratungsdateien` WHERE `parent` = '{termin}'""".format(termin=termin.beratung_name), as_dict=True)[0].qty > 0 else 0
        else:
            termin.eintrittsdatum = None
            termin.creation = None
            termin.hat_attachement = None
    
    berater_in = get_berater_in_from_hash(berater_in)
    email = get_email_from_berater_in(berater_in)
    html_von = frappe.utils.getdate(von).strftime("%d.%m.%Y") if von else ''
    html_bis = frappe.utils.getdate(bis).strftime("%d.%m.%Y") if bis else ''
    html = frappe.render_template("mvd/mvd/page/individueller_arbeitsplan/pdf.html", {'berater_in': berater_in, 'email': email, 'termine': termine, 'von': html_von, 'bis': html_bis})
    from frappe.utils.pdf import get_pdf
    pdf = get_pdf(html)
    datum = frappe.utils.getdate(von or None).strftime("%Y-%m-%d")
    wochentag = _(frappe.utils.getdate(von or None).strftime("%A"), "de")[:2].upper()
    clean_berater_name = "unbekannt"
    if berater_in:
        clean_berater_name = berater_in.split(" ")[0].replace("ä", "ae").replace("ö", "oe").replace("ü", "ue")
    frappe.local.response.filename = "{datum}_{wochentag}_{name}.pdf".format(datum=datum, wochentag=wochentag, name=clean_berater_name)
    frappe.local.response.filecontent = pdf
    frappe.local.response.type = "download"


@frappe.whitelist()
def get_arbeitsplan_word(berater_in, von=None, bis=None):
    termine = get_termin_uebersicht(berater_in, von, bis)
    for termin in termine:
        new_von_zeit = "{0}:{1}".format("{0}".format(termin.von).split(" ")[1].split(":")[0], "{0}".format(termin.von).split(" ")[1].split(":")[1])
        new_von = "{0}.{1}.{2}".format("{0}".format(termin.von).split(" ")[0].split("-")[2], "{0}".format(termin.von).split(" ")[0].split("-")[1], "{0}".format(termin.von).split(" ")[0].split("-")[0])
        new_bis = "{0}:{1}".format("{0}".format(termin.bis).split(" ")[1].split(":")[0], "{0}".format(termin.bis).split(" ")[1].split(":")[1])
        termin.von_zeit = new_von_zeit
        termin.von = new_von
        termin.bis = new_bis
        if termin.abp_referenz:
            termin.eintrittsdatum = frappe.utils.getdate(termin.eintrittsdatum).strftime("%d.%m.%Y")
            termin.creation = frappe.utils.getdate(termin.creation).strftime("%d.%m.%Y")
            termin.hat_attachement = 1 if frappe.db.sql("""SELECT COUNT(`name`) AS `qty` FROM `tabBeratungsdateien` WHERE `parent` = '{termin}'""".format(termin=termin.beratung_name), as_dict=True)[0].qty > 0 else 0
        else:
            termin.eintrittsdatum = None
            termin.creation = None
            termin.hat_attachement = None
    
    
    berater_in = get_berater_in_from_hash(berater_in)
    email = get_email_from_berater_in(berater_in)
    html_von = frappe.utils.getdate(von).strftime("%d.%m.%Y") if von else ''
    html_bis = frappe.utils.getdate(bis).strftime("%d.%m.%Y") if bis else ''
    datum = frappe.utils.getdate(von or None).strftime("%Y-%m-%d")
    wochentag = _(frappe.utils.getdate(von or None).strftime("%A"), "de")[:2].upper()
    clean_berater_name = "unbekannt"
    if berater_in:
        clean_berater_name = berater_in.split(" ")[0].replace("ä", "ae").replace("ö", "oe").replace("ü", "ue")

    word_dict = {
        'berater_in': berater_in,
        'email': email,
        'von': html_von,
        'bis': html_bis,
        'termine': termine,
        'file_name': "{datum}_{wochentag}_{name}.docx".format(datum=datum, wochentag=wochentag, name=clean_berater_name)
    }

    create_word_doc(word_dict)

    return "/files/{0}".format(word_dict['file_name'])

def create_word_doc(word_dict):
    from docx import Document
    word_doc = Document()
    word_doc.add_heading("{0} - {1}".format(word_dict['berater_in'], word_dict['email']), level=1)
    if word_dict['von'] and word_dict['bis']:
        word_doc.add_paragraph("({0} - {1})".format(word_dict['von'], word_dict['bis']))
    
    termin_loop = 1
    for termin in word_dict['termine']:
        if "abp_referenz" in termin:
            # Create Table
            table = word_doc.add_table(rows=4, cols=4)

            # Table Header 1
            cell = table.cell(0, 0)
            cell.text = ""
            p = cell.add_paragraph()
            run = p.add_run("Wann")
            run.bold = True
            cell = table.cell(0, 1)
            cell.text = ""
            p = cell.add_paragraph()
            run = p.add_run("Wo")
            run.bold = True
            cell = table.cell(0, 2)
            cell.text = ""
            p = cell.add_paragraph()
            run = p.add_run("Wie")
            run.bold = True
            cell = table.cell(0, 3)
            cell.text = ""
            p = cell.add_paragraph()
            run = p.add_run("Telefonnummer")
            run.bold = True

            # Table Content 1
            table.cell(1, 0).text = "{0}\n{1} - {2}".format(termin['von'], termin['von_zeit'], termin['bis'])
            table.cell(1, 1).text = termin['ort'].replace("(" + termin['sektion_id'] + ")", "")
            table.cell(1, 2).text = termin['art']
            table.cell(1, 3).text = "{0}".format(termin['telefonnummer'] or "---")

            # Table Header 2
            cell1 = table.cell(2, 0)
            cell2 = table.cell(2, 1)
            merged_cell = cell1.merge(cell2)
            merged_cell.text = ""
            p = merged_cell.add_paragraph()
            run = p.add_run("Mitglied")
            run.bold = True
            cell1 = table.cell(2, 2)
            cell2 = table.cell(2, 3)
            merged_cell = cell1.merge(cell2)
            merged_cell.text = ""
            p = merged_cell.add_paragraph()
            run = p.add_run("Details")
            run.bold = True

            # Table Content 2
            cell1 = table.cell(3, 0)
            cell2 = table.cell(3, 1)
            merged_cell = cell1.merge(cell2)
            merged_cell.text = "{0}\n{1}".format(termin['mitglied_nr'], termin['adressblock'].replace("<br>", "\n"))
            cell1 = table.cell(3, 2)
            cell2 = table.cell(3, 3)
            merged_cell = cell1.merge(cell2)
            merged_cell.text = "{0} {1} {2}\n\nEintritt: {3}\nBez. Jahr: {4}\nTermin vereinbart am: {5} {6}".format(termin.beratungskategorie.split(" - ")[0] if termin.beratungskategorie else '____', \
                                                termin.beratungskategorie_2.split(" - ")[0] if termin.beratungskategorie_2 else '____', \
                                                termin.beratungskategorie_3.split(" - ")[0] if termin.beratungskategorie_3 else '____', \
                                                termin.eintrittsdatum if termin.eintrittsdatum else '---', \
                                                termin.bezahltes_mitgliedschaftsjahr, termin.creation, termin.created_by)
            
            word_doc.add_paragraph("\nUnterlagen > {0}\nBeratung : ☐ wurde durchgeführt ☐ Person nicht erschienen\nNotiz/Bemerkungen: {1}".format("Ja" if termin.hat_attachement == 1 else "Nein", \
                                                                                                                                                  termin.notiz or "---"))
            
            if termin_loop == 2:
                word_doc.add_page_break()
                termin_loop = 1
            else:
                termin_loop += 1
    physical_path = "/home/frappe/frappe-bench/sites/{0}/public/files/{1}".format(frappe.local.site_path.replace("./", ""), word_dict['file_name'])
    word_doc.save(physical_path)

def get_berater_in_from_hash(hash):
    berater_in = frappe.db.sql("""SELECT `name` FROM `tabTermin Kontaktperson` WHERE `md_hash` = '{0}'""".format(hash), as_dict=True)
    if len(berater_in) > 0:
        return berater_in[0].name
    else:
        return None
    
def get_email_from_berater_in(berater_in):
    """
    Returns the email address of the first berater_in as str
    """
    email = frappe.db.sql("""SELECT `user` 
                    FROM `tabTermin Kontaktperson Multi User` 
                    WHERE `parent` = '{berater_in}';"""
                                            .format(berater_in=berater_in), as_dict=True)
    if len(email) > 0:
        return email[0].user
    else:
        return "keine E-Mail"

def get_created_by(owner):
    vor_nachname = frappe.db.sql("""SELECT `first_name`, `last_name` FROM `tabUser` WHERE `name` = '{0}'""".format(owner), as_dict=True)
    if len(vor_nachname) < 1:
        return '/??'
    
    kuerzel = '/{0}{1}'.format(vor_nachname[0].first_name[0] if vor_nachname[0].first_name else '?', vor_nachname[0].last_name[0] if vor_nachname[0].last_name else '?')
    return kuerzel